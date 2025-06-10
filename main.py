import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_LINE_SPACING,WD_COLOR_INDEX,WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import io
#input_path = "Change.docx"
#output_path = "Reformatted_Change.docx"
st.title("DOCX Formatter")
uploaded_file = st.file_uploader("Upload your Word (.docx) file", type=["docx"])



if uploaded_file:
# Load document
    doc = Document(uploaded_file)
    formatted_doc = Document()
    
    sections = doc.sections
    formatted_sections = formatted_doc.sections
    for fsection in formatted_doc.sections:
        fsection.top_margin = Inches(1)
        fsection.bottom_margin = Inches(1)
        fsection.left_margin = Inches(1)
        fsection.right_margin = Inches(1)
        
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    # for i, section in enumerate(doc.sections):
    #     top = section.top_margin.inches
    #     bottom = section.bottom_margin.inches
    #     left = section.left_margin.inches
    #     right = section.right_margin.inches
    #     print(top)
    #     print(bottom)
    #     print(left)
    #     print(right)
    #     section.top_margin = Inches(1)
    #     section.bottom_margin = Inches(1)
    #     section.left_margin = Inches(1)
    #     section.right_margin = Inches(1)
    def force_font_on_paragraph(paragraph, font_name="Times New Roman", size=12):
        run = paragraph.add_run(" ")
        font = run.font
        font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        font.size = Pt(size)
        
        # Force paragraph style manually if needed
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)

    def insert_page_numbers_except_first(document):
        for section in document.sections:
            # Enable different first page
            section.different_first_page_header_footer = True

            # Get footer for all pages except first
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Add only the page number field { PAGE }
            run = paragraph.add_run()

            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = 'PAGE'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')

            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run._r.append(fldChar3)
    def clean_text(text):
        text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces/newlines with a single space
        return text.strip()
    
    def clean_tabulated_text(text):
    # Удаляем табуляции и лишние пробелы
        text = re.sub(r'\s+', ' ', text)  # заменяет все виды пробелов, включая табы, на один пробел
        text = re.sub(r'\s+([.,:;!?])', r'\1', text)  # удаляет пробел перед знаками препинания
        text = re.sub(r'([.,:;!?])([^\s])', r'\1 \2', text)  # ставит пробел после знаков препинания, если его нет
        text = re.sub(r'\s{2,}', ' ', text)  # двойные пробелы
        return text.strip()

    def clean_text_extended(text):
    # Remove invisible characters and carriage returns
        text = re.sub(r'[\u200B-\u200D\uFEFF]', '', text)
        
        # Replace all newlines, tabs, and multiple spaces with a single space
        text = re.sub(r'[\n\r\t]+', ' ', text)
        text = re.sub(r'\s{2,}', ' ', text)
        # [\n\r\t]+: removes any newlines (\n), carriage returns (\r), and tabs (\t).

        # \s{2,}: collapses double/multiple spaces into one.

        # strip(): trims leading/trailing whitespace.
        text = text.replace('\n','')
        return text.strip()
    def clean_from_invisible_char(text):
        text = re.sub()
        text = text.replace('\r', '')
    def shade_paragraph(paragraph, color="D9D9D9"):
        p_pr = paragraph._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        p_pr.append(shd)
    def apply_format(paragraph, font_size, bold, align, spacing_after=3, spacing_before=0):
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        font = run.font
        font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        font.size = Pt(font_size)
        font.bold = bold
        font.color.rgb = RGBColor(0, 0, 0)
        pf = paragraph.paragraph_format
        pf.alignment = align
        pf.space_after = Pt(spacing_after)
        pf.space_before = Pt(spacing_before)
        
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE  # <--- обязательно добавьте эту строку
        pf.line_spacing = 1.0  
       # pf.line_spacing = Pt(12)

    def apply_format2(paragraph, font_size, bold, align, spacing_after, spacing_before=0):
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        font = run.font
        font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        font.size = Pt(font_size)
        font.bold = bold
        font.color.rgb = RGBColor(0, 0, 0)
        pf = paragraph.paragraph_format
        pf.alignment = align
        pf.space_after = Pt(spacing_after)
        pf.space_before = Pt(spacing_before)
        #pf.line_spacing = Pt(12)

    def set_cell_shading(cell, color):
    # Add shading to a table cell
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)  # Hex color, e.g., 'D9D9D9'
        tcPr.append(shd)
    def move_short_words_to_next_line(text):
        lines = text.split('\n')
        new_lines = []
        
        for line in lines:
            words = line.strip().split()
            if len(words) >= 2 and len(words[-1]) > 2:
                if len(words[-2]) <= 2:
                    # move short word to next line
                    new_line = " ".join(words[:-2])
                    short_word = words[-2]
                    last_word = words[-1]
                    if new_line:
                        new_lines.append(new_line)
                    new_lines.append(f"{short_word}\u00A0{last_word}")
                    continue
            new_lines.append(line)
        
        return '\n'.join(new_lines)

    def apply_typographic_fixes(text):
        # 1. Non-breaking space after №
        # text = text.replace("№ ", "№\u00A0")  
        # text = text.replace("№", "№\u00A0")
        # text = re.sub(r'№\s*(\d+)', r'№\u00A0\1', text)
        nbsp = "\u00A0"  # non-breaking space
        text = re.sub(r'№\s*(\d+)', f'№{nbsp}\\1', text)
        # 2. Non-breaking space between day and month (using raw string pattern, regular string replacement)
        text = re.sub(
            r'\b(\d{1,2})\s+(января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря)\b',
            lambda m: f"{m.group(1)}\u00A0{m.group(2)}",
            text
        )

        # 3. Move 1–2 letter word from end of line to start of next
        #text = move_short_words_to_next_line(text)

        return text

    def set_character_spacing(run, spacing_val):
        rPr = run._element.get_or_add_rPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:val'), str(spacing_val))
        rPr.append(spacing)

    def fix_docx_numbering(text):
        """
        Fix duplicated numbering at the beginning of a line: '1. 1. Text' -> '1. Text'
        """
        return re.sub(r'^(\d+\.)\s+\d+\.\s+', r'\1 ', text)

    def strip_manual_numbering(text):
        return re.sub(r'^\s*(\d+[\.\)]|[-–]\s*\d+[\.\)]?)\s*', '', text)
    
    def set_format(paragraph, size=12, bold=False, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(size)
        font.bold = bold
        font.color.rgb = RGBColor(0, 0, 0)
        paragraph.paragraph_format.alignment = align

    def add_signature_table(title, name):
        table = formatted_doc.add_table(rows=1, cols=2)
        table.columns[0].width = Inches(4.7)
        table.columns[1].width = Inches(1.8)
        cell1, cell2 = table.rows[0].cells
        p1 = cell1.paragraphs[0].add_run(title)
        p1.bold = True
        p1.font.size = Pt(12)
        p1.font.name = 'Times New Roman'
        p2 = cell2.paragraphs[0].add_run(name)
        p2.font.size = Pt(12)
        p2.font.name = 'Times New Roman'
        cell2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    def set_shading(run, fill_color):
        """Applies background shading (highlighting) to a run."""
        rPr = run._element.get_or_add_rPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)  # e.g., 'FFFF00' for yellow
        rPr.append(shd)

    def clear_headers_and_footers(doc: Document):
        for section in doc.sections:
            # Очистка верхнего колонтитула
            header = section.header
            for para in header.paragraphs:
                para.clear()

            # Очистка нижнего колонтитула
            footer = section.footer
            for para in footer.paragraphs:
                para.clear()
    def bold_keywords(paragraph, keywords):
        for keyword in keywords:
            if keyword in paragraph.text:
                # Store existing text and clear paragraph
                text = paragraph.text
                paragraph.clear()
                i = 0
                while i < len(text):
                    matched = False
                    for word in keywords:
                        if text[i:].startswith(word):
                            run = paragraph.add_run(word)
                            run.bold = True
                            run.font.name = 'Times New Roman'
                            i += len(word)
                            matched = True
                            break
                    if not matched:
                        run = paragraph.add_run(text[i])
                        i += 1
    
    cleaned_paragraphs = [
        clean_text_extended(para.text)
        for para in doc.paragraphs
        if clean_text_extended(para.text)  # filters out empty ones
    ]
    # for i, para in enumerate(cleaned_paragraphs):
    #     print(f"{i+1}. {para}")
        
    # === Split paragraphs by blocks ===
    blocks = {}
    current_block = None
    
    clear_headers_and_footers(doc)
    for para in doc.paragraphs:
        text = para.text.strip()
        #text = clean_text_extended(para.text.strip())
        text = clean_text_extended(text)
        if not text:
            continue
        #print(repr(para.text))
        #print(para[0].text)
        #print(text)
        block_header = re.match(r"^Блок(\d+)", text)
        #block_header = {}
        if block_header:
            current_block = f"Блок{block_header.group(1)}"
            blocks[current_block] = []
        elif current_block:
            blocks[current_block].append(text)
    #insert_page_numbers_except_first(doc)
    # === Process blocks with specific styles ===
    if doc.paragraphs and doc.paragraphs[0].text.strip() == "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА":
        for block, paragraphs in blocks.items():
            #formatted_doc.add_paragraph(block).runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red block header

            if block == "Блок1":
                text = "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА"
                p = formatted_doc.add_paragraph()
                run = p.add_run(text)
                set_character_spacing(run,60)
                apply_format(p,14,True,WD_PARAGRAPH_ALIGNMENT.CENTER)
                pf = p.paragraph_format
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                pf.line_spacing = 1.0
                pf.space_before = Pt(0)
                pf.space_after = Pt(6)
                
                for para in paragraphs:
                    para = clean_text(para)
                    #para = para.replace("№", "№\u00A0")
                    p = formatted_doc.add_paragraph(para)
                    p.text = apply_typographic_fixes(p.text)
                    set_format(p, 12, True, WD_PARAGRAPH_ALIGNMENT.CENTER)
                    
                    pf = p.paragraph_format
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    pf.line_spacing = 1.0
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
            elif block == "Блок2":
                # add 
                emptyLine = formatted_doc.add_paragraph()
                # run = emptyLine.add_run(" ")
                # font = run.font
                # font.name = 'Times New Roman'
                # run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                # font.size = Pt(12)
                force_font_on_paragraph(emptyLine)
                
                parEmptyLine = emptyLine.paragraph_format
                parEmptyLine.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                parEmptyLine.line_spacing = 1.0
                parEmptyLine.space_before = Pt(0)
                parEmptyLine.space_after = Pt(0)
                
                text = "Условные (сокращенные) обозначения, использованные в пояснительной записке"
                p = formatted_doc.add_paragraph()
                run = p.add_run(text)
                apply_format(p,11,True,WD_PARAGRAPH_ALIGNMENT.LEFT)
                for para in paragraphs:
                    para = clean_text(para)
                    if "–" in para or "-" in para:
                        p = formatted_doc.add_paragraph(para)
                        set_format(p, 11, False)
                        p.text = apply_typographic_fixes(p.text)
                        apply_format(p,11,False,WD_PARAGRAPH_ALIGNMENT.LEFT)
                    else:
                        parts = para.split("–", 1)
                        if len(parts) == 2:
                            term, desc = parts
                            p = formatted_doc.add_paragraph()
                            p.text = apply_typographic_fixes(p.text)
                            run = p.add_run(term.strip() + " – ")
                            run.bold = True
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            p.add_run(desc.strip())
                            set_format(p, 11)
                            apply_format(p,11,False,WD_PARAGRAPH_ALIGNMENT.LEFT)
                emptyLine = formatted_doc.add_paragraph()
                # run = emptyLine.add_run(" ")
                # font = run.font
                # font.name = 'Times New Roman'
                # run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                # font.size = Pt(12)
                force_font_on_paragraph(emptyLine)
                
                parEmptyLine = emptyLine.paragraph_format
                parEmptyLine.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                parEmptyLine.line_spacing = 1.0
                parEmptyLine.space_before = Pt(0)
                parEmptyLine.space_after = Pt(0)
                #formatted_doc.add_paragraph()
            # elif block == "Блок3":
            #     #lines_to_merge = []
            #     #for p in doc.paragraphs:
            #     #    if p.text.strip():  # skip empty lines
            #     #        lines_to_merge.append(p.text)

            #     for para in paragraphs:
            #         p = formatted_doc.add_paragraph(para)
            #         run = p.add_run()
            #         run = p.runs[0] if p.runs else p.add_run()
            #         font = run.font
            #         font.name = 'Times New Roman'
            #         run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            #         font.size = Pt(11)
            #         font.bold = False
            #         font.color.rgb = RGBColor(0, 0, 0)
            #         set_shading(run, 'D9D9D9')
            #     # Join all lines with a line break
            #     #for idx, line in enumerate(lines_to_merge):
            #     #    run.add_text(line)
            #     #    if idx != len(lines_to_merge) - 1:
            #     #        run.add_break() 

            #         pf = p.paragraph_format
            #         pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            #         pf.space_after = Pt(3)
            #         pf.space_before = Pt(0)
            #         pf.line_spacing = 1.0  
            #         pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            #         if p.text == "Основание выноса вопроса на рассмотрение Советом директоров":
            #             font.bold = True
            elif block == "Блок3":
                
                # table = formatted_doc.add_table(rows=1, cols=1)
                # cell = table.cell(0, 0)
                # set_cell_shading(cell, 'D9D9D9')

                # for para in paragraphs:
                #     p = cell.add_paragraph(para)
                #     run = p.runs[0] if p.runs else p.add_run()
                #     font = run.font
                #     font.name = 'Times New Roman'
                #     run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                #     font.size = Pt(11)
                #     font.bold = False
                #     font.color.rgb = RGBColor(0, 0, 0)

                #     pf = p.paragraph_format
                #     pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                #     pf.space_after = Pt(3)
                #     pf.space_before = Pt(0)
                #     pf.line_spacing = 1.0
                #     pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

                #     if para.strip() == "Основание выноса вопроса на рассмотрение Советом директоров":
                #         font.bold = True
                for para in paragraphs: 
                    #formatted_doc.add_paragraph()
                    para = clean_text(para)
                    p = formatted_doc.add_paragraph(para)
                    run = p.add_run()
                    run = p.runs[0] if p.runs else p.add_run()
                    font = run.font

                    font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    font.size = Pt(11)
                    font.bold = False
                    # if para.strip() == "Основание выноса вопроса на рассмотрение Советом директоров":
                    #     font.bold = True
                    if para.strip().startswith("Основание"):
                        font.bold = True
                    pf = p.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(2)  # Уменьшает расстояние после строки
                    pf.line_spacing = 1.0 
                    shade_paragraph(p)
                    
            elif block == "Блок4":
                emptyLine = formatted_doc.add_paragraph()
                # run = emptyLine.add_run(" ")
                # font = run.font
                # font.name = 'Times New Roman'
                # run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                # font.size = Pt(12)
                force_font_on_paragraph(emptyLine)
                
                parEmptyLine = emptyLine.paragraph_format
                parEmptyLine.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                parEmptyLine.line_spacing = 1.0
                parEmptyLine.space_before = Pt(0)
                parEmptyLine.space_after = Pt(0)
                
                for para in paragraphs: 
                    para = clean_text(para)
                    p = formatted_doc.add_paragraph(para)
                    p.text = apply_typographic_fixes(p.text)
                    run = p.add_run()
                    run = p.runs[0] if p.runs else p.add_run()
                    font = run.font
                    font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    font.size = Pt(12)
                    font.bold = False
                    font.color.rgb = RGBColor(0, 0, 0)

                    pf = p.paragraph_format
                    pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    pf.space_after = Pt(0)
                    pf.space_before = Pt(0)
                    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    #set_shading(run, 'D9D9D9')
                emptyLine = formatted_doc.add_paragraph()
                # run = emptyLine.add_run(" ")
                # font = run.font
                # font.name = 'Times New Roman'
                # run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                # font.size = Pt(12)
                force_font_on_paragraph(emptyLine)
                
                parEmptyLine = emptyLine.paragraph_format
                parEmptyLine.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                parEmptyLine.line_spacing = 1.0
                parEmptyLine.space_before = Pt(0)
                parEmptyLine.space_after = Pt(0)
            # elif block == "Блок5":
            #     for para in paragraphs: 
                        
            #             fixed_text = apply_typographic_fixes(para) 
            #             p = formatted_doc.add_paragraph(fixed_text)
                        
            #             #if para.startswith(""): 

            #             pf = p.paragraph_format
            #             pf.first_line_indent = Inches(0.3)                     # Выступ (отступ первой строки)
            #             pf.space_before = Pt(0)                                # Интервал перед абзацем
            #             pf.space_after = Pt(3)                                 # Интервал после абзаца
            #             pf.line_spacing = 1.0                                  # Одинарный межстрочный
            #             pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            #             pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY       
            #             #set_format(p)
            #             #apply_typographic_fixes(p.text)
            #             run = p.add_run()
            #             run = p.runs[0] if p.runs else p.add_run()
            #             font = run.font
            #             font.name = 'Times New Roman'
            #             run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            #             font.size = Pt(11)
            # elif block == "Блок5":
            #     for para in paragraphs:
            #         fixed_text = apply_typographic_fixes(para)
                    
            #         # Проверка наличия тире
            #         if "–" in fixed_text:
            #             parts = fixed_text.split("–", 1)
            #             term, desc = parts[0].strip(), parts[1].strip()

            #             p = formatted_doc.add_paragraph()
            #             pf = p.paragraph_format
            #             pf.first_line_indent = Inches(0.3)
            #             pf.space_before = Pt(0)
            #             pf.space_after = Pt(3)
            #             pf.line_spacing = 1.0
            #             pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            #             pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            #             # Первая часть (жирная)
            #             run = p.add_run(term + " –\t")
            #             run.font.bold = True
            #             run.font.name = 'Times New Roman'
            #             run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            #             run.font.size = Pt(11)

            #             # Вторая часть (обычная)
            #             run2 = p.add_run(desc)
            #             run2.font.name = 'Times New Roman'
            #             run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            #             run2.font.size = Pt(11)

            #         else:
            #             # Обычный абзац
            #             p = formatted_doc.add_paragraph(fixed_text)
            #             pf = p.paragraph_format
            #             pf.first_line_indent = Inches(0.3)
            #             pf.space_before = Pt(0)
            #             pf.space_after = Pt(3)
            #             pf.line_spacing = 1.0
            #             pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            #             pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            #             run = p.runs[0]
            #             run.font.name = 'Times New Roman'
            #             run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            #             run.font.size = Pt(11)
            # elif block == "Блок5":
            #     for para in paragraphs: 
            #         fixed_text = apply_typographic_fixes(para) 

            #         # Создаем абзац с текстом
            #         p = formatted_doc.add_paragraph(fixed_text)

            #         # Форматирование абзаца
            #         pf = p.paragraph_format
            #         pf.first_line_indent = Inches(0.3)                 # Отступ первой строки (выступ)
            #         pf.space_before = Pt(0)                            # Интервал перед абзацем
            #         pf.space_after = Pt(3)                             # Интервал после абзаца
            #         pf.line_spacing = 1.0                              # Одинарный межстрочный интервал
            #         pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            #         pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY      # Выровнять по ширине

            #         # Безопасный доступ к run
            #         if p.runs:
            #             run = p.runs[0]
            #         else:
            #             run = p.add_run()
            #         # Настройка шрифта
            #         run.font.name = 'Times New Roman'
            #         run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            #         run.font.size = Pt(11)
            elif block == "Блок5":
                if paragraphs:
                    first_para = formatted_doc.add_paragraph()
                    #first_para = clean_text(first_para)
                    pf = first_para.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(3)
                    pf.line_spacing = 1.0
                    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                    run = first_para.add_run(paragraphs[0].strip())
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    run.font.size = Pt(11)
                    shade_paragraph(first_para)
                for para in paragraphs[1:]:
        # Проверка на наличие тире
                    para = clean_text(para)
                    para = apply_typographic_fixes(para)
                    match = re.match(r"^[–\-]\s*(.*)", para.strip())
                    if match:
                            desc = match.group(1)
                            # Создаем абзац
                            p = formatted_doc.add_paragraph()
                            pf = p.paragraph_format

                            # Висячий отступ
                            pf.left_indent = Inches(0.3)
                            pf.first_line_indent = Inches(-0.3)
                            pf.tab_stops.add_tab_stop(Inches(0.3), WD_TAB_ALIGNMENT.LEFT)
                            # Прочее форматирование
                            pf.space_before = Pt(0)
                            pf.space_after = Pt(3)
                            pf.line_spacing = 1.0
                            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                            # Добавляем tab stop (необязательно, но полезно для ручного редактирования)
                            #pf.tab_stops.add_tab_stop(Inches(0.3), WD_TAB_ALIGNMENT.LEFT)

                            # Добавляем тире, табуляцию и текст
                            run = p.add_run("–\t" + desc.strip())
                            run.font.name = 'Times New Roman'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                            run.font.size = Pt(11)
                            #run.font.highlight_color = WD_COLOR_INDEX.GRAY_50
                            shade_paragraph(p)
                    else:
                        # Абзац без тире — просто отформатированный
                        p = formatted_doc.add_paragraph()
                        pf = p.paragraph_format
                        
                        font = p.add_run().font
                        font.highlight_color = WD_COLOR_INDEX.GRAY_50
                        
                        # Такой же висячий отступ для выравнивания
                        pf.left_indent = Inches(0.3)
                        pf.first_line_indent = Inches(-0.3)
                        #pf.left_indent = Inches(0.3)
                        #pf.first_line_indent = Inches(-0.3)
                        pf.tab_stops.add_tab_stop(Inches(0.3), WD_TAB_ALIGNMENT.LEFT)
                        
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(3)
                        pf.line_spacing = 1.0
                        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        #pf.tab_stops.add_tab_stop(Inches(0.3), WD_TAB_ALIGNMENT.LEFT)
                        run = p.add_run("\t" + para.strip())
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        run.font.size = Pt(11)
                        #run.font.highlight_color = WD_COLOR_INDEX.GRAY_50
                        shade_paragraph(p)

            elif block == "Блок6":
                for para in paragraphs: 
                    # para = clean_text(para)
                    # para = apply_typographic_fixes(para)
                    
                    # ps = formatted_doc.add_paragraph(pa)
                    
                    # psf = ps.paragraph_format
                    # psf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    # psf.space_before = Pt(0)
                    # psf.space_after = Pt(6)
                    # psf.line_spacing = 1.0
                    # psf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            

                    # runf = ps.runs[0] if ps.runs else ps.add_run()
                    # runf.font.name = 'Times New Roman'
                    # runf._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    # runf.font.size = Pt(12)
                    para = clean_text(para)
                    para = apply_typographic_fixes(para)

                    ps = formatted_doc.add_paragraph()  # <-- empty paragraph
                    runf = ps.add_run(para)             # <-- add run manually
                    runf.font.name = 'Times New Roman'
                    runf._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    runf.font.size = Pt(12)

                    psf = ps.paragraph_format
                    psf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    psf.space_before = Pt(0)
                    psf.space_after = Pt(6)
                    psf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    psf.line_spacing = 1.0
                
            elif block == "Блок7":
                # first_nonempty = ""
                # for para in paragraphs:
                #     clean_para = para.strip().replace(" ", "").replace("*", "*")
                #     if clean_para:
                #         first_nonempty = para.strip()
                #         break

                # # Normalize and compare
                # normalized = re.sub(r"\s+", "", first_nonempty)
                num = 0
                # if normalized != "* * *":
                for para in paragraphs: 
                    if para[0] != "* * *":
                        if num == 0:
                            p = formatted_doc.add_paragraph("* * *")
                            
                        psf = p.paragraph_format
                        psf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        psf.space_before = Pt(0)
                        psf.space_after = Pt(6)
                        psf.spacing = 1.0
                        psf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            
                        runf = p.runs[0]
                        runf.font.name = 'Times New Roman'
                        runf._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        runf.font.size = Pt(12)
                        
                        num += 1
                    
                for idx, para in enumerate(paragraphs):
                    para = clean_text(para)
                    para = apply_typographic_fixes(para)

                    if idx == 0:
                        para = para.replace("→", "").strip()  # Убираем символ стрелки
                        p = formatted_doc.add_paragraph()
                        pf = p.paragraph_format
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(3)
                        pf.line_spacing = 1.0
                        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        run = p.add_run(para)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        run.font.size = Pt(12)
                        continue

                    # Обработка пунктов "1.", "2." и т.д.
                    if re.match(r"^\d+\.", para.strip()):
                        numbered_text = fix_docx_numbering(para.strip())
                        numbered_text = re.sub(r"^(\d+\.)\s*", r"\1\t", numbered_text)
                        p = formatted_doc.add_paragraph()
                        pf = p.paragraph_format
                        pf.left_indent = Inches(0.5)
                        pf.first_line_indent = Inches(-0.5)
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(3)
                        pf.line_spacing = 1.0
                        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        run = p.add_run(numbered_text)
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        run.font.size = Pt(12)
                        shade_paragraph(p)
                    else:
                        # Прочий текст — отступ и табуляция
                        p = formatted_doc.add_paragraph()
                        pf = p.paragraph_format
                        pf.left_indent = Inches(0.5)
                        pf.first_line_indent = Inches(-0.5)
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(3)
                        pf.line_spacing = 1.0
                        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        run = p.add_run("\t" + para.strip())
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        run.font.size = Pt(12)
                        shade_paragraph(p)
                        
            elif block == "Блок8":
            #------------------------------------------------
                bold_words = [
                    "Председатель Правления", 
                    "Заместитель председателя Правления", 
                    "Советник председателя Правления", 
                    "Управляющий директор", 
                ]
                for _ in range(5):
                    emptyLine = formatted_doc.add_paragraph()
                    force_font_on_paragraph(emptyLine)

                    parEmptyLine = emptyLine.paragraph_format
                    parEmptyLine.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    parEmptyLine.line_spacing = 1.0
                    parEmptyLine.space_before = Pt(0)
                    parEmptyLine.space_after = Pt(0)
                    
                for para in paragraphs:
                    print(para)
                    para = para.strip()
                    matched_role = None

                    # Check if paragraph starts with one of the roles

                    for bw in bold_words:
                        if para.startswith(bw):
                            matched_role = bw   
                            break
                        
                    if matched_role:
                        name = para[len(matched_role):].strip()  # extract the rest as name
                        p = formatted_doc.add_paragraph()
                        p.text = apply_typographic_fixes(p.text)
                        # Set tab stop
                        tab_stops = p.paragraph_format.tab_stops
                        tab_stops.add_tab_stop(Inches(4.5), alignment=WD_TAB_ALIGNMENT.LEFT)

                        # Add role (bold)
                        run_role = p.add_run(matched_role)
                        run_role.font.name = "Times New Roman"
                        run_role.font.size = Pt(12)
                        run_role.bold = True

                        # Add tab + name (bold)
                        p.add_run("\t")
                        run_name = p.add_run(name)
                        run_name.font.name = "Times New Roman"
                        run_name._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        run_name.font.size = Pt(12)
                        run_name.bold = True

                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        print("here matched")
                    else:
                        # No role matched — just output as is
                        p = formatted_doc.add_paragraph(para)
                        p.text = apply_typographic_fixes(p.text)
                        run = p.runs[0] if p.runs else p.add_run()
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
                        run.bold = True
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        
                        print("here not matched")
                    # else:
                    #     # Add paragraph manually with formatting
                    #     clean_text2 = apply_typographic_fixes(para)
                    #     p = formatted_doc.add_paragraph()
                    #     run = p.add_run(clean_text2)
                    #     run.font.name = "Times New Roman"
                    #     run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    #     run.font.size = Pt(10)
                    #     run.bold = True
                    #     p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    #     print("here not matched")
                        
            elif block == "Блок9":
                emptyLine = formatted_doc.add_paragraph()
                force_font_on_paragraph(emptyLine)

                parEmptyLine = emptyLine.paragraph_format
                parEmptyLine.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                parEmptyLine.line_spacing = 1.0
                parEmptyLine.space_before = Pt(0)
                parEmptyLine.space_after = Pt(0)
                
                for para in paragraphs: 
                    para = clean_text(para)
                    if para:
                        p = formatted_doc.add_paragraph()
                        p.text = apply_typographic_fixes(p.text)
                        run = p.add_run(para)
                        #apply_format(p,10,False,WD_PARAGRAPH_ALIGNMENT.LEFT)
                        font = run.font
                        font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        font.size = Pt(10)
                
                emptyLine = formatted_doc.add_paragraph()
                run = emptyLine.add_run(" ")  # Add a space to ensure formatting applies
                font = run.font
                font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                font.size = Pt(12)
                force_font_on_paragraph(emptyLine)

                parEmptyLine = emptyLine.paragraph_format
                parEmptyLine.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                parEmptyLine.line_spacing = 1.0
                parEmptyLine.space_before = Pt(0)
                parEmptyLine.space_after = Pt(0)
                
            elif block == "Блок10":
                cleaned_paragraphs = [clean_text(p) for p in paragraphs if clean_text(p)]

# Check if "Приложение" already exists as the first non-empty line
                if not cleaned_paragraphs or cleaned_paragraphs[0].strip().lower() != "приложение":

                    p = formatted_doc.add_paragraph()
                    run = p.add_run("Приложения")
                    font = run.font
                    font.name = 'Times New Roman'
                    font.size = Pt(11)
                    font.bold = True
                    
                    pf = p.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after =  Pt(3)
                for para in cleaned_paragraphs:
                    if para.strip().lower() == "приложения":
                        continue
                    para = clean_text(para)
                    para = strip_manual_numbering(para)
                    if para:
                        p = formatted_doc.add_paragraph(para, style='List Number') 
                        p.text = apply_typographic_fixes(p.text)
                        run = p.runs[0] if p.runs else p.add_run()
                        font = run.font
                        font.name = 'Times New Roman'
                        font.size = Pt(11)
                        font.bold = False
                        font.color.rgb = RGBColor(0, 0, 0)
                        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        
            else:
                for para in paragraphs:
                    para = clean_text(para)
                    if para:
                        fixed_text = apply_typographic_fixes(para) 
                        p = formatted_doc.add_paragraph(para)
                        set_format(p)

    if doc.paragraphs and doc.paragraphs[0].text.strip() == "БЮЛЛИТЕНЬ":
        
        
        
        print("here")
    insert_page_numbers_except_first(formatted_doc)
    insert_page_numbers_except_first(doc)
    buffer = io.BytesIO()
    formatted_doc.save(buffer)
    buffer.seek(0)
    st.download_button("Download Reformatted DOCX", data=buffer, file_name=uploaded_file.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
